from prompt import PROMPT
import io
from docx import Document
from docx.shared import Inches
import streamlit as st
from langchain_google_genai import GoogleGenerativeAI
from langchain.prompts import PromptTemplate


# # Api key setup
# load_dotenv()
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Api key setup
GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]


# Response Generation function
def generate_proposal(company_name, company_industry, company_website,client_name, client_industry, client_website,
                      proposal_type, prompt_addition):

        # Initialize LLM Model
        llm = GoogleGenerativeAI(model="gemini-2.0-flash",api_key=GOOGLE_API_KEY)

        # Prepare the prompt template with required variables
        prompt = PromptTemplate(
            input_variables=[
                "company_name", "company_industry", "company_website","client_name", "client_industry", "client_website",
                "proposal_type", "prompt_addition"
            ],
            template=PROMPT
        )

        # Build LLMChain
        # llm_chain = LLMChain(llm=llm, prompt=prompt)
        llm_chain = prompt | llm

        # Run the chain with provided inputs
        response = llm_chain.invoke({
            "company_name": company_name,
            "company_industry": company_industry,
            "company_website": company_website,
            "client_name": client_name,
            "client_industry": client_industry,
            "client_website": client_website,
            "proposal_type": proposal_type,
            "prompt_addition": prompt_addition
        })

        return response


# Function to save proposal and image to a DOCX file
def save_proposal_to_docx(company_name, proposal_response, logo_file):
    # Create a new Word document
    doc = Document()

    # Add company name as title
    doc.add_heading(f"{company_name} - Proposal", 0)

    # Add logo image if uploaded
    if logo_file is not None:
        image_stream = io.BytesIO(logo_file.read())
        doc.add_picture(image_stream, width=Inches(2))
        logo_file.seek(0)  

    # Add proposal content
    doc.add_heading("Generated Proposal", level=1)
    doc.add_paragraph(proposal_response)

    # Save document
    filename = f"{company_name}_proposal.docx"
    doc.save(filename)

    return filename


# Set page config
st.set_page_config(page_title="Expert Proposal Generator", page_icon="📄",layout="wide")

# Title
st.title("📄 Expert Proposal Generator")

# Sidebar Inputs
st.sidebar.title("📑 Proposal Inputs")

with st.sidebar.form("proposal_form"):
    st.subheader("🔹 Company Details")
    company_name = st.text_input("Company Name")
    company_industry = st.text_input("Industry")
    company_website = st.text_input("Website Link")

    st.subheader("🔹 Client Details")
    client_name = st.text_input("Client Name")
    client_industry = st.text_input("Industry", key="client_industry")
    client_website = st.text_input("Website Link", key="client_website")

    st.subheader("🔹 Proposal Type & Prompt")
    proposal_type = st.text_input("Type of Proposal")
    prompt_addition = st.text_area("Additional Prompt / Instructions")

    st.subheader("🔹 Company Logo")
    logo_file = st.file_uploader("Upload Logo (optional)", type=["png", "jpg", "jpeg"])

    submitted = st.form_submit_button("Generate Proposal")



if submitted:
    with st.spinner("⏳ Generating your proposal..."):
        # Generate the proposal using the provided inputs
        proposal_response = generate_proposal(company_name, company_industry, company_website, client_name, client_industry,
                                              client_website, proposal_type, prompt_addition)

    # Layout with 2 columns: proposal text and logo
    col1, col2 = st.columns([5, 1])

    with col2:
        if logo_file is not None:
            st.image(logo_file, width=120)  

    with col1:
        st.subheader("📄 Generated Proposal")
        st.write(proposal_response)

    
    # Save proposal to DOCX file
    filename = save_proposal_to_docx(company_name, proposal_response, logo_file)

    # Provide download link
    with open(filename, "rb") as file:
        st.download_button(
            label="📥 Download Proposal Document",
            data=file,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


