# docx_generator.py
import re
from docx import Document
from docx.shared import Pt
from io import BytesIO

def generate_proposal_docx_data(generate_section_content_func, SECTION_GUIDELINES, firm, client_name, product, proposal_date):
    """
    Generate a proposal DOCX document.
    The function collects the proposal sections using the provided generate_section_content_func,
    adds bullet points per sentence, and returns the document as a BytesIO stream.
    """
    doc = Document()
    # Header information
    header = doc.add_paragraph()
    header_run = header.add_run(
        f"Consultancy Firm: {firm}\nClient: {client_name}\nProduct: {product}\nProposal Date: {proposal_date}\n"
    )
    header_run.bold = True
    doc.add_heading(f"{product} Proposal", level=0)
    
    final_sections = [
        "Executive Summary",
        "About Us",
        "Understanding the Client’s Needs",
        "Scope of Work",
        "Deliverables",
        "Timeline",
        "Pricing and Payment Terms",
        "Client Responsibilities",
        "Confidentiality & Compliance",
        "Terms & Conditions",
        "Acceptance of Proposal"
    ]
    
    for section in final_sections:
        guideline = SECTION_GUIDELINES.get(
            section,
            "Provide a concise summary without any chain-of-thought or filler language."
        )
        doc.add_heading(section, level=1)
        section_content = generate_section_content_func(
            section, firm, client_name, product, proposal_date, guideline
        )
        # Split content into sentences and add as bullet points
        lines = section_content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            sentences = re.split(r'\.\s+', line)
            for sentence in sentences:
                sentence = sentence.strip()
                if sentence:
                    if not sentence.endswith('.'):
                        sentence += '.'
                    doc.add_paragraph(sentence, style='List Bullet')
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

if __name__ == "__main__":
    # Dummy function for testing standalone
    def dummy_generate_section_content(section, firm, client_name, product, proposal_date, guideline):
        return f"Dummy content for {section}."
    
    # For standalone testing, you can uncomment and run:
    # generate_proposal_docx_data(dummy_generate_section_content, {}, "Test Firm", "Test Client", "Test Product", "01-01-2025")
